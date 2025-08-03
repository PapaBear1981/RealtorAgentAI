"""
Multi-format output generation for contracts.

This module provides comprehensive output generation capabilities
supporting PDF, DOCX, HTML, and TXT formats with proper styling
and formatting for professional contract documents.
"""

import logging
import io
import tempfile
import os
from datetime import datetime
from typing import Dict, Any, Optional, BinaryIO
from pathlib import Path

# PDF generation
try:
    import weasyprint
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logging.warning("WeasyPrint not available - PDF generation disabled")

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX generation disabled")

# HTML processing
from bs4 import BeautifulSoup
import re

from ..models.template import OutputFormat
from ..core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class OutputGenerationError(Exception):
    """Exception raised during output generation."""
    pass


class OutputGenerator:
    """
    Multi-format output generator for contracts.
    
    Supports PDF, DOCX, HTML, and TXT output formats with
    professional styling and formatting.
    """
    
    def __init__(self):
        """Initialize output generator."""
        self.default_css = self._get_default_css()
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> list[OutputFormat]:
        """Get list of supported output formats."""
        formats = [OutputFormat.HTML, OutputFormat.TXT]
        
        if WEASYPRINT_AVAILABLE:
            formats.append(OutputFormat.PDF)
        
        if DOCX_AVAILABLE:
            formats.append(OutputFormat.DOCX)
        
        return formats
    
    def _get_default_css(self) -> str:
        """Get default CSS for contract styling."""
        return """
        @page {
            size: letter;
            margin: 1in;
            @bottom-center {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
                color: #666;
            }
        }
        
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #000;
            margin: 0;
            padding: 0;
        }
        
        .contract-header {
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #000;
            padding-bottom: 20px;
        }
        
        .contract-header h1 {
            font-size: 18pt;
            font-weight: bold;
            margin: 0 0 10px 0;
            text-transform: uppercase;
        }
        
        .contract-header .version {
            font-size: 10pt;
            color: #666;
            margin: 5px 0;
        }
        
        .contract-section {
            margin-bottom: 25px;
        }
        
        .contract-section h2 {
            font-size: 14pt;
            font-weight: bold;
            margin: 20px 0 10px 0;
            border-bottom: 1px solid #ccc;
            padding-bottom: 5px;
        }
        
        .contract-section h3 {
            font-size: 12pt;
            font-weight: bold;
            margin: 15px 0 8px 0;
        }
        
        .contract-clause {
            margin-bottom: 15px;
            text-align: justify;
        }
        
        .contract-clause.numbered {
            counter-increment: clause;
        }
        
        .contract-clause.numbered::before {
            content: counter(clause) ". ";
            font-weight: bold;
        }
        
        .signature-section {
            margin-top: 50px;
            page-break-inside: avoid;
        }
        
        .signature-line {
            border-bottom: 1px solid #000;
            width: 300px;
            display: inline-block;
            margin: 20px 10px 5px 0;
        }
        
        .signature-label {
            font-size: 10pt;
            color: #666;
            margin-top: 5px;
        }
        
        .signature-block {
            display: inline-block;
            vertical-align: top;
            margin: 20px 30px 20px 0;
            min-width: 250px;
        }
        
        .date-line {
            border-bottom: 1px solid #000;
            width: 150px;
            display: inline-block;
            margin: 20px 10px 5px 0;
        }
        
        .terms-list {
            counter-reset: term;
        }
        
        .terms-list li {
            counter-increment: term;
            margin-bottom: 10px;
        }
        
        .terms-list li::marker {
            content: "(" counter(term, lower-alpha) ") ";
        }
        
        .highlight {
            background-color: #ffff99;
            padding: 2px 4px;
        }
        
        .important {
            font-weight: bold;
            color: #d00;
        }
        
        .small-text {
            font-size: 10pt;
            color: #666;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        table th, table td {
            border: 1px solid #000;
            padding: 8px;
            text-align: left;
        }
        
        table th {
            background-color: #f0f0f0;
            font-weight: bold;
        }
        
        .page-break {
            page-break-before: always;
        }
        
        .no-break {
            page-break-inside: avoid;
        }
        """
    
    async def generate_output(
        self,
        content: str,
        output_format: OutputFormat,
        title: str = "Contract Document",
        custom_css: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Generate output in specified format.
        
        Args:
            content: HTML content to convert
            output_format: Target output format
            title: Document title
            custom_css: Custom CSS styles
            metadata: Document metadata
            
        Returns:
            bytes: Generated document content
        """
        try:
            if output_format not in self.supported_formats:
                raise OutputGenerationError(f"Format {output_format} not supported")
            
            # Clean and prepare content
            content = self._prepare_content(content, title, custom_css)
            
            if output_format == OutputFormat.HTML:
                return await self._generate_html(content, title, metadata)
            elif output_format == OutputFormat.PDF:
                return await self._generate_pdf(content, title, metadata)
            elif output_format == OutputFormat.DOCX:
                return await self._generate_docx(content, title, metadata)
            elif output_format == OutputFormat.TXT:
                return await self._generate_txt(content, title, metadata)
            else:
                raise OutputGenerationError(f"Unsupported format: {output_format}")
                
        except Exception as e:
            logger.error(f"Output generation failed: {e}")
            raise OutputGenerationError(f"Failed to generate {output_format} output: {str(e)}")
    
    def _prepare_content(self, content: str, title: str, custom_css: Optional[str] = None) -> str:
        """Prepare and clean HTML content."""
        # Parse HTML
        soup = BeautifulSoup(content, 'html.parser')
        
        # Add document structure if missing
        if not soup.find('html'):
            # Wrap content in proper HTML structure
            css_content = self.default_css
            if custom_css:
                css_content += "\n" + custom_css
            
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>{css_content}</style>
</head>
<body>
{content}
</body>
</html>"""
            return html_content
        
        # Add CSS if not present
        if not soup.find('style') and not soup.find('link', rel='stylesheet'):
            style_tag = soup.new_tag('style')
            css_content = self.default_css
            if custom_css:
                css_content += "\n" + custom_css
            style_tag.string = css_content
            
            head = soup.find('head')
            if head:
                head.append(style_tag)
        
        return str(soup)
    
    async def _generate_html(
        self, 
        content: str, 
        title: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate HTML output."""
        return content.encode('utf-8')
    
    async def _generate_pdf(
        self, 
        content: str, 
        title: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate PDF output using WeasyPrint."""
        if not WEASYPRINT_AVAILABLE:
            raise OutputGenerationError("PDF generation not available - WeasyPrint not installed")
        
        try:
            # Create HTML document
            html_doc = HTML(string=content)
            
            # Generate PDF
            pdf_bytes = html_doc.write_pdf()
            
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise OutputGenerationError(f"PDF generation failed: {str(e)}")
    
    async def _generate_docx(
        self, 
        content: str, 
        title: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate DOCX output using python-docx."""
        if not DOCX_AVAILABLE:
            raise OutputGenerationError("DOCX generation not available - python-docx not installed")
        
        try:
            # Parse HTML content
            soup = BeautifulSoup(content, 'html.parser')
            
            # Create new document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = title
            doc.core_properties.author = "Contract Generation System"
            doc.core_properties.created = datetime.utcnow()
            
            if metadata:
                doc.core_properties.subject = metadata.get('subject', '')
                doc.core_properties.keywords = metadata.get('keywords', '')
            
            # Configure styles
            self._configure_docx_styles(doc)
            
            # Convert HTML to DOCX
            self._html_to_docx(soup, doc)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.read()
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise OutputGenerationError(f"DOCX generation failed: {str(e)}")
    
    async def _generate_txt(
        self, 
        content: str, 
        title: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate plain text output."""
        try:
            # Parse HTML and extract text
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Add title header
            if title:
                text = f"{title}\n{'=' * len(title)}\n\n{text}"
            
            return text.encode('utf-8')
            
        except Exception as e:
            logger.error(f"TXT generation failed: {e}")
            raise OutputGenerationError(f"TXT generation failed: {str(e)}")
    
    def _configure_docx_styles(self, doc):
        """Configure DOCX document styles."""
        try:
            styles = doc.styles
            
            # Configure Normal style
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(12)
            
            # Create contract header style
            if 'Contract Header' not in [s.name for s in styles]:
                header_style = styles.add_style('Contract Header', WD_STYLE_TYPE.PARAGRAPH)
                header_font = header_style.font
                header_font.name = 'Times New Roman'
                header_font.size = Pt(18)
                header_font.bold = True
                header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_style.paragraph_format.space_after = Pt(20)
            
        except Exception as e:
            logger.warning(f"Failed to configure DOCX styles: {e}")
    
    def _html_to_docx(self, soup, doc):
        """Convert HTML soup to DOCX document."""
        try:
            # Find body content
            body = soup.find('body')
            if not body:
                body = soup
            
            # Process elements
            for element in body.children:
                if hasattr(element, 'name'):
                    self._process_html_element(element, doc)
                elif element.strip():
                    # Text node
                    doc.add_paragraph(element.strip())
                    
        except Exception as e:
            logger.warning(f"HTML to DOCX conversion error: {e}")
            # Fallback: add all text as paragraphs
            text = soup.get_text()
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
    
    def _process_html_element(self, element, doc):
        """Process individual HTML element for DOCX conversion."""
        try:
            tag_name = element.name.lower()
            
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Headers
                level = int(tag_name[1])
                text = element.get_text().strip()
                if text:
                    if level == 1:
                        para = doc.add_paragraph(text, style='Contract Header')
                    else:
                        para = doc.add_paragraph(text, style='Heading ' + str(min(level, 3)))
            
            elif tag_name == 'p':
                # Paragraphs
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            
            elif tag_name in ['ul', 'ol']:
                # Lists
                for li in element.find_all('li'):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet' if tag_name == 'ul' else 'List Number')
            
            elif tag_name == 'table':
                # Tables
                self._process_table(element, doc)
            
            elif tag_name == 'div':
                # Divs - process children
                for child in element.children:
                    if hasattr(child, 'name'):
                        self._process_html_element(child, doc)
                    elif child.strip():
                        doc.add_paragraph(child.strip())
            
            else:
                # Other elements - extract text
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)
                    
        except Exception as e:
            logger.warning(f"Error processing HTML element {element.name}: {e}")
    
    def _process_table(self, table_element, doc):
        """Process HTML table for DOCX."""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return
            
            # Count columns
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            
            # Create table
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            
            # Fill table
            for row_idx, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols:
                        table.cell(row_idx, col_idx).text = cell.get_text().strip()
                        
        except Exception as e:
            logger.warning(f"Error processing table: {e}")


# Global output generator instance
_output_generator: Optional[OutputGenerator] = None


def get_output_generator() -> OutputGenerator:
    """
    Get global output generator instance.
    
    Returns:
        OutputGenerator: Configured output generator
    """
    global _output_generator
    
    if _output_generator is None:
        _output_generator = OutputGenerator()
    
    return _output_generator


# Export generator
__all__ = [
    "OutputGenerator",
    "OutputGenerationError",
    "get_output_generator",
]
