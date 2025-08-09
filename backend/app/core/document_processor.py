"""
Document processing module for PDF, DOCX, and OCR operations.

This module provides comprehensive document processing capabilities
including text extraction, OCR, and metadata extraction.
"""

import asyncio
import io
import logging
from datetime import datetime
from typing import Dict, Any, Optional, List, BinaryIO, Tuple
from pathlib import Path
import tempfile
import os

# Document processing imports
# import pymupdf4llm  # Temporarily disabled for debugging
# import fitz  # PyMuPDF  # Temporarily disabled for debugging
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from .config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Base exception for document processing operations."""
    pass


class DocumentProcessor:
    """
    Comprehensive document processor for various file types.

    Supports PDF text extraction, DOCX processing, OCR for images,
    and metadata extraction with async processing capabilities.
    """

    def __init__(self):
        """Initialize document processor."""
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_docx,  # Will attempt to process as DOCX
            'txt': self._process_text,
            'image': self._process_image_ocr
        }

        # Configure Tesseract if path is provided
        if hasattr(settings, 'TESSERACT_CMD_PATH') and settings.TESSERACT_CMD_PATH:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD_PATH
        elif hasattr(settings, 'TESSERACT_CMD') and settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    async def process_document(
        self,
        file_content: bytes,
        file_type: str,
        filename: str,
        processing_options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process document and extract text and metadata.

        Args:
            file_content: Raw file content as bytes
            file_type: Type of file (pdf, docx, image, etc.)
            filename: Original filename
            processing_options: Optional processing parameters

        Returns:
            Dict: Processing results with extracted text and metadata
        """
        try:
            logger.info(f"Processing document: {filename} (type: {file_type})")

            # Get processing function
            processor_func = self.supported_formats.get(file_type)
            if not processor_func:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")

            # Set default processing options
            options = processing_options or {}

            # Process document in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                processor_func,
                file_content,
                filename,
                options
            )

            # Add processing metadata
            result['processing_metadata'] = {
                'processed_at': datetime.utcnow().isoformat(),
                'file_type': file_type,
                'filename': filename,
                'processor_version': '1.0.0',
                'processing_options': options
            }

            logger.info(f"Successfully processed document: {filename}")
            return result

        except Exception as e:
            logger.error(f"Failed to process document {filename}: {e}")
            raise DocumentProcessingError(f"Document processing failed: {e}")

    def _process_pdf(
        self,
        file_content: bytes,
        filename: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process PDF document using PyMuPDF.

        Args:
            file_content: PDF file content
            filename: Original filename
            options: Processing options

        Returns:
            Dict: Extracted text and metadata
        """
        try:
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name

            try:
                # Use pymupdf4llm for advanced text extraction
                if False:  # options.get('use_advanced_extraction', True):  # Temporarily disabled for debugging
                    # Extract text with markdown formatting
                    # markdown_text = pymupdf4llm.to_markdown(temp_path)

                    # Also get page chunks for detailed analysis
                    # page_chunks = pymupdf4llm.to_markdown(temp_path, page_chunks=True)

                    # extracted_text = markdown_text
                    # pages_data = page_chunks
                    pass
                else:
                    # Use basic PyMuPDF extraction
                    doc = fitz.open(temp_path)
                    extracted_text = ""
                    pages_data = []

                    for page_num in range(doc.page_count):
                        page = doc[page_num]
                        page_text = page.get_text()
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}"

                        pages_data.append({
                            'page': page_num + 1,
                            'text': page_text,
                            'metadata': {
                                'rotation': page.rotation,
                                'rect': list(page.rect),
                                'mediabox': list(page.mediabox)
                            }
                        })

                    doc.close()

                # Extract document metadata using PyMuPDF
                doc = fitz.open(temp_path)
                metadata = doc.metadata

                # Get document statistics
                stats = {
                    'page_count': doc.page_count,
                    'has_links': any(page.get_links() for page in doc),
                    'has_images': any(page.get_images() for page in doc),
                    'has_annotations': any(page.annots() for page in doc),
                    'is_encrypted': doc.needs_pass,
                    'file_size': len(file_content)
                }

                doc.close()

                # Perform OCR if requested and document has images
                ocr_text = None
                if options.get('perform_ocr', False) and stats['has_images']:
                    try:
                        ocr_text = self._extract_pdf_images_ocr(temp_path)
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for {filename}: {ocr_error}")

                return {
                    'extracted_text': extracted_text,
                    'ocr_text': ocr_text,
                    'pages_data': pages_data,
                    'document_metadata': metadata,
                    'document_stats': stats,
                    'processing_success': True
                }

            finally:
                # Clean up temporary file
                os.unlink(temp_path)

        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise DocumentProcessingError(f"PDF processing error: {e}")

    def _process_docx(
        self,
        file_content: bytes,
        filename: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process DOCX document using python-docx.

        Args:
            file_content: DOCX file content
            filename: Original filename
            options: Processing options

        Returns:
            Dict: Extracted text and metadata
        """
        try:
            # Load document from bytes
            doc_stream = io.BytesIO(file_content)
            doc = DocxDocument(doc_stream)

            # Extract text from paragraphs
            paragraphs = []
            full_text = ""

            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    paragraphs.append({
                        'text': para_text,
                        'style': para.style.name if para.style else None
                    })
                    full_text += para_text + "\n"

            # Extract text from tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                        full_text += cell_text + " "
                    table_data.append(row_data)
                    full_text += "\n"
                tables_data.append(table_data)

            # Extract document properties
            props = doc.core_properties
            metadata = {
                'title': props.title,
                'author': props.author,
                'subject': props.subject,
                'keywords': props.keywords,
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
                'last_modified_by': props.last_modified_by,
                'revision': props.revision,
                'category': props.category,
                'comments': props.comments
            }

            # Document statistics
            stats = {
                'paragraph_count': len(paragraphs),
                'table_count': len(tables_data),
                'character_count': len(full_text),
                'word_count': len(full_text.split()),
                'file_size': len(file_content)
            }

            return {
                'extracted_text': full_text.strip(),
                'ocr_text': None,  # DOCX doesn't need OCR
                'paragraphs': paragraphs,
                'tables': tables_data,
                'document_metadata': metadata,
                'document_stats': stats,
                'processing_success': True
            }

        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise DocumentProcessingError(f"DOCX processing error: {e}")

    def _process_text(
        self,
        file_content: bytes,
        filename: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process plain text file.

        Args:
            file_content: Text file content
            filename: Original filename
            options: Processing options

        Returns:
            Dict: Text content and metadata
        """
        try:
            # Decode text with fallback encoding detection
            try:
                text_content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text_content = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    text_content = file_content.decode('utf-8', errors='replace')

            # Basic text statistics
            lines = text_content.split('\n')
            stats = {
                'line_count': len(lines),
                'character_count': len(text_content),
                'word_count': len(text_content.split()),
                'file_size': len(file_content)
            }

            return {
                'extracted_text': text_content,
                'ocr_text': None,
                'lines': lines,
                'document_metadata': {'encoding': 'utf-8'},
                'document_stats': stats,
                'processing_success': True
            }

        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise DocumentProcessingError(f"Text processing error: {e}")

    def _process_image_ocr(
        self,
        file_content: bytes,
        filename: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process image file using OCR.

        Args:
            file_content: Image file content
            filename: Original filename
            options: Processing options

        Returns:
            Dict: OCR extracted text and metadata
        """
        try:
            # Load image
            image = Image.open(io.BytesIO(file_content))

            # Get image metadata
            image_metadata = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'has_transparency': image.mode in ('RGBA', 'LA') or 'transparency' in image.info
            }

            # Perform OCR
            ocr_config = options.get('ocr_config', '--oem 3 --psm 6')
            ocr_language = options.get('ocr_language', 'eng')

            # Extract text using Tesseract
            ocr_text = pytesseract.image_to_string(
                image,
                lang=ocr_language,
                config=ocr_config
            )

            # Get detailed OCR data if requested
            ocr_data = None
            if options.get('detailed_ocr', False):
                ocr_data = pytesseract.image_to_data(
                    image,
                    lang=ocr_language,
                    config=ocr_config,
                    output_type=pytesseract.Output.DICT
                )

            # Image statistics
            stats = {
                'width': image.size[0],
                'height': image.size[1],
                'channels': len(image.getbands()),
                'file_size': len(file_content),
                'ocr_confidence': self._calculate_ocr_confidence(ocr_data) if ocr_data else None
            }

            return {
                'extracted_text': None,  # Images don't have native text
                'ocr_text': ocr_text.strip(),
                'ocr_data': ocr_data,
                'image_metadata': image_metadata,
                'document_stats': stats,
                'processing_success': True
            }

        except Exception as e:
            logger.error(f"Image OCR processing failed: {e}")
            raise DocumentProcessingError(f"Image OCR error: {e}")

    def _extract_pdf_images_ocr(self, pdf_path: str) -> str:
        """
        Extract images from PDF and perform OCR.

        Args:
            pdf_path: Path to PDF file

        Returns:
            str: OCR text from all images
        """
        try:
            doc = fitz.open(pdf_path)
            ocr_texts = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                image_list = page.get_images()

                for img_index, img in enumerate(image_list):
                    try:
                        # Extract image
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)

                        # Convert to PIL Image
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("ppm")
                            pil_image = Image.open(io.BytesIO(img_data))

                            # Perform OCR
                            ocr_text = pytesseract.image_to_string(pil_image)
                            if ocr_text.strip():
                                ocr_texts.append(f"--- Page {page_num + 1}, Image {img_index + 1} ---\n{ocr_text}")

                        pix = None  # Free memory

                    except Exception as img_error:
                        logger.warning(f"Failed to process image {img_index} on page {page_num}: {img_error}")
                        continue

            doc.close()
            return "\n\n".join(ocr_texts)

        except Exception as e:
            logger.error(f"PDF image OCR failed: {e}")
            return ""

    def _calculate_ocr_confidence(self, ocr_data: Dict[str, Any]) -> float:
        """
        Calculate average OCR confidence score.

        Args:
            ocr_data: OCR data from Tesseract

        Returns:
            float: Average confidence score
        """
        if not ocr_data or 'conf' not in ocr_data:
            return 0.0

        confidences = [conf for conf in ocr_data['conf'] if conf > 0]
        return sum(confidences) / len(confidences) if confidences else 0.0

    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported file formats.

        Returns:
            List[str]: Supported formats
        """
        return list(self.supported_formats.keys())

    async def validate_document(
        self,
        file_content: bytes,
        file_type: str,
        filename: str
    ) -> Dict[str, Any]:
        """
        Validate document without full processing.

        Args:
            file_content: File content to validate
            file_type: Expected file type
            filename: Original filename

        Returns:
            Dict: Validation results
        """
        try:
            validation_result = {
                'is_valid': False,
                'file_type': file_type,
                'detected_type': None,
                'errors': [],
                'warnings': []
            }

            # Basic file validation
            if len(file_content) == 0:
                validation_result['errors'].append("File is empty")
                return validation_result

            # Type-specific validation
            if file_type == 'pdf':
                try:
                    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                        temp_file.write(file_content)
                        temp_path = temp_file.name

                    try:
                        doc = fitz.open(temp_path)
                        validation_result['detected_type'] = 'pdf'
                        validation_result['is_valid'] = True

                        if doc.needs_pass:
                            validation_result['warnings'].append("PDF is password protected")

                        doc.close()
                    finally:
                        os.unlink(temp_path)

                except Exception as e:
                    validation_result['errors'].append(f"Invalid PDF: {e}")

            elif file_type in ['docx', 'doc']:
                try:
                    doc_stream = io.BytesIO(file_content)
                    doc = DocxDocument(doc_stream)
                    validation_result['detected_type'] = 'docx'
                    validation_result['is_valid'] = True
                except Exception as e:
                    validation_result['errors'].append(f"Invalid DOCX: {e}")

            elif file_type == 'image':
                try:
                    image = Image.open(io.BytesIO(file_content))
                    validation_result['detected_type'] = 'image'
                    validation_result['is_valid'] = True

                    # Check image size
                    if image.size[0] * image.size[1] > 50000000:  # 50MP limit
                        validation_result['warnings'].append("Very large image may slow processing")

                except Exception as e:
                    validation_result['errors'].append(f"Invalid image: {e}")

            elif file_type == 'txt':
                try:
                    file_content.decode('utf-8')
                    validation_result['detected_type'] = 'txt'
                    validation_result['is_valid'] = True
                except UnicodeDecodeError:
                    validation_result['warnings'].append("Text encoding may cause issues")
                    validation_result['detected_type'] = 'txt'
                    validation_result['is_valid'] = True

            return validation_result

        except Exception as e:
            logger.error(f"Document validation failed: {e}")
            return {
                'is_valid': False,
                'file_type': file_type,
                'detected_type': None,
                'errors': [f"Validation error: {e}"],
                'warnings': []
            }


# Global document processor instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """
    Get global document processor instance.

    Returns:
        DocumentProcessor: Configured document processor
    """
    global _document_processor

    if _document_processor is None:
        _document_processor = DocumentProcessor()

    return _document_processor


# Export classes and functions
__all__ = [
    "DocumentProcessingError",
    "DocumentProcessor",
    "get_document_processor",
]
