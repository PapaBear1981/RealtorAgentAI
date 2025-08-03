"""
Export tasks for document generation and delivery.

This module contains Celery tasks for generating PDF/DOCX documents,
preparing deliveries, and sending notifications.
"""

import logging
import tempfile
from typing import Dict, Any, Optional, List
from datetime import datetime
from pathlib import Path
from io import BytesIO

from celery import current_task
import structlog
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Inches

from ..core.celery_app import celery_app, DatabaseTask
from ..core.storage import get_storage_client
from ..models.contract import Contract
from ..models.audit_log import AuditLog, AuditAction

logger = structlog.get_logger(__name__)


@celery_app.task(bind=True, base=DatabaseTask, name="app.tasks.export_tasks.generate_pdf_document")
def generate_pdf_document(
    self,
    contract_id: int,
    user_id: int,
    template_options: Optional[Dict[str, Any]] = None,
    include_metadata: bool = True
) -> Dict[str, Any]:
    """
    Generate PDF document from contract content.
    
    Args:
        contract_id: Database contract record ID
        user_id: User requesting the generation
        template_options: PDF formatting options
        include_metadata: Include contract metadata in PDF
        
    Returns:
        Dict: Generation results with storage information
    """
    try:
        logger.info(
            "Starting PDF document generation",
            contract_id=contract_id,
            user_id=user_id,
            include_metadata=include_metadata
        )
        
        # Get contract record
        contract = self.session.get(Contract, contract_id)
        if not contract:
            raise ValueError(f"Contract record not found: {contract_id}")
        
        # Default template options
        options = template_options or {}
        page_size = options.get("page_size", letter)
        font_size = options.get("font_size", 12)
        margins = options.get("margins", {"top": 72, "bottom": 72, "left": 72, "right": 72})
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            topMargin=margins["top"],
            bottomMargin=margins["bottom"],
            leftMargin=margins["left"],
            rightMargin=margins["right"]
        )
        
        # Build PDF content
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title = f"Real Estate Contract - {contract.title or 'Untitled'}"
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        
        # Contract metadata
        if include_metadata and contract.metadata:
            story.append(Paragraph("Contract Information", styles['Heading2']))
            
            metadata_info = [
                f"Contract ID: {contract.id}",
                f"Created: {contract.created_at.strftime('%Y-%m-%d %H:%M:%S') if contract.created_at else 'Unknown'}",
                f"Status: {contract.status}",
                f"Deal ID: {contract.deal_id}" if contract.deal_id else None,
                f"Template ID: {contract.template_id}" if contract.template_id else None
            ]
            
            for info in metadata_info:
                if info:
                    story.append(Paragraph(info, styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        # Contract content
        story.append(Paragraph("Contract Content", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        # Split content into paragraphs and add to PDF
        content = contract.content or "No content available"
        paragraphs = content.split('\n\n')
        
        for paragraph in paragraphs:
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Build PDF
        start_time = datetime.utcnow()
        doc.build(story)
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        
        # Get PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        # Store PDF in storage
        storage_client = get_storage_client()
        filename = f"contract_{contract_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
        storage_key = storage_client.generate_storage_key(filename, "pdf", user_id)
        
        upload_result = storage_client.upload_file(
            BytesIO(pdf_content),
            storage_key,
            "application/pdf",
            {"contract_id": str(contract_id), "generated_by": str(user_id)}
        )
        
        # Create audit log
        audit_log = AuditLog(
            user_id=user_id,
            actor=f"system:export_task:{self.request.id}",
            action=AuditAction.CONTRACT_EXPORT,
            resource_type="contract",
            resource_id=str(contract_id),
            success=True,
            meta={
                "contract_id": contract_id,
                "export_format": "pdf",
                "file_size": len(pdf_content),
                "storage_key": storage_key,
                "processing_time": processing_time
            }
        )
        self.session.add(audit_log)
        self.session.commit()
        
        results = {
            "status": "completed",
            "contract_id": contract_id,
            "format": "pdf",
            "storage_key": storage_key,
            "file_size": len(pdf_content),
            "filename": filename,
            "metadata": {
                "processing_time": processing_time,
                "page_count": len(story),
                "generation_timestamp": datetime.utcnow().isoformat(),
                "task_id": self.request.id
            }
        }
        
        logger.info(
            "PDF document generation completed",
            contract_id=contract_id,
            file_size=len(pdf_content),
            processing_time=processing_time
        )
        
        return results
        
    except Exception as exc:
        logger.error(
            "PDF document generation failed",
            contract_id=contract_id,
            error=str(exc),
            exc_info=True
        )
        
        # Create error audit log
        audit_log = AuditLog(
            user_id=user_id,
            actor=f"system:export_task:{self.request.id}",
            action=AuditAction.CONTRACT_EXPORT,
            resource_type="contract",
            resource_id=str(contract_id),
            success=False,
            error_message=str(exc),
            meta={"contract_id": contract_id, "export_format": "pdf"}
        )
        self.session.add(audit_log)
        self.session.commit()
        
        # Retry with exponential backoff
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@celery_app.task(bind=True, base=DatabaseTask, name="app.tasks.export_tasks.generate_docx_document")
def generate_docx_document(
    self,
    contract_id: int,
    user_id: int,
    template_options: Optional[Dict[str, Any]] = None,
    include_metadata: bool = True
) -> Dict[str, Any]:
    """
    Generate DOCX document from contract content.
    
    Args:
        contract_id: Database contract record ID
        user_id: User requesting the generation
        template_options: DOCX formatting options
        include_metadata: Include contract metadata in document
        
    Returns:
        Dict: Generation results with storage information
    """
    try:
        logger.info(
            "Starting DOCX document generation",
            contract_id=contract_id,
            user_id=user_id,
            include_metadata=include_metadata
        )
        
        # Get contract record
        contract = self.session.get(Contract, contract_id)
        if not contract:
            raise ValueError(f"Contract record not found: {contract_id}")
        
        # Create DOCX document
        doc = Document()
        
        # Add title
        title = f"Real Estate Contract - {contract.title or 'Untitled'}"
        doc.add_heading(title, 0)
        
        # Add contract metadata
        if include_metadata and contract.metadata:
            doc.add_heading('Contract Information', level=1)
            
            metadata_info = [
                f"Contract ID: {contract.id}",
                f"Created: {contract.created_at.strftime('%Y-%m-%d %H:%M:%S') if contract.created_at else 'Unknown'}",
                f"Status: {contract.status}",
                f"Deal ID: {contract.deal_id}" if contract.deal_id else None,
                f"Template ID: {contract.template_id}" if contract.template_id else None
            ]
            
            for info in metadata_info:
                if info:
                    doc.add_paragraph(info)
        
        # Add contract content
        doc.add_heading('Contract Content', level=1)
        
        content = contract.content or "No content available"
        paragraphs = content.split('\n\n')
        
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to memory buffer
        start_time = datetime.utcnow()
        buffer = BytesIO()
        doc.save(buffer)
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        
        # Get DOCX content
        docx_content = buffer.getvalue()
        buffer.close()
        
        # Store DOCX in storage
        storage_client = get_storage_client()
        filename = f"contract_{contract_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        storage_key = storage_client.generate_storage_key(filename, "docx", user_id)
        
        upload_result = storage_client.upload_file(
            BytesIO(docx_content),
            storage_key,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            {"contract_id": str(contract_id), "generated_by": str(user_id)}
        )
        
        # Create audit log
        audit_log = AuditLog(
            user_id=user_id,
            actor=f"system:export_task:{self.request.id}",
            action=AuditAction.CONTRACT_EXPORT,
            resource_type="contract",
            resource_id=str(contract_id),
            success=True,
            meta={
                "contract_id": contract_id,
                "export_format": "docx",
                "file_size": len(docx_content),
                "storage_key": storage_key,
                "processing_time": processing_time
            }
        )
        self.session.add(audit_log)
        self.session.commit()
        
        results = {
            "status": "completed",
            "contract_id": contract_id,
            "format": "docx",
            "storage_key": storage_key,
            "file_size": len(docx_content),
            "filename": filename,
            "metadata": {
                "processing_time": processing_time,
                "generation_timestamp": datetime.utcnow().isoformat(),
                "task_id": self.request.id
            }
        }
        
        logger.info(
            "DOCX document generation completed",
            contract_id=contract_id,
            file_size=len(docx_content),
            processing_time=processing_time
        )
        
        return results
        
    except Exception as exc:
        logger.error(
            "DOCX document generation failed",
            contract_id=contract_id,
            error=str(exc),
            exc_info=True
        )
        
        # Create error audit log
        audit_log = AuditLog(
            user_id=user_id,
            actor=f"system:export_task:{self.request.id}",
            action=AuditAction.CONTRACT_EXPORT,
            resource_type="contract",
            resource_id=str(contract_id),
            success=False,
            error_message=str(exc),
            meta={"contract_id": contract_id, "export_format": "docx"}
        )
        self.session.add(audit_log)
        self.session.commit()
        
        # Retry with exponential backoff
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@celery_app.task(bind=True, name="app.tasks.export_tasks.prepare_document_delivery")
def prepare_document_delivery(
    self,
    storage_key: str,
    recipient_email: str,
    delivery_method: str = "email",
    expiration_hours: int = 24
) -> Dict[str, Any]:
    """
    Prepare document for delivery to recipient.
    
    Args:
        storage_key: Storage location of the document
        recipient_email: Email address of recipient
        delivery_method: Delivery method (email, download_link)
        expiration_hours: Hours until delivery link expires
        
    Returns:
        Dict: Delivery preparation results
    """
    try:
        logger.info(
            "Preparing document delivery",
            storage_key=storage_key,
            recipient_email=recipient_email,
            delivery_method=delivery_method
        )
        
        # Generate secure delivery link
        from secrets import token_urlsafe
        delivery_token = token_urlsafe(32)
        
        # Calculate expiration
        expiration_time = datetime.utcnow() + timedelta(hours=expiration_hours)
        
        # Store delivery information (would be in database in production)
        delivery_info = {
            "storage_key": storage_key,
            "recipient_email": recipient_email,
            "delivery_token": delivery_token,
            "delivery_method": delivery_method,
            "created_at": datetime.utcnow().isoformat(),
            "expires_at": expiration_time.isoformat(),
            "status": "prepared"
        }
        
        # Generate delivery URL
        from ..core.config import get_settings
        settings = get_settings()
        base_url = getattr(settings, 'BASE_URL', 'http://localhost:8000')
        delivery_url = f"{base_url}/api/documents/delivery/{delivery_token}"
        
        results = {
            "status": "prepared",
            "delivery_token": delivery_token,
            "delivery_url": delivery_url,
            "recipient_email": recipient_email,
            "delivery_method": delivery_method,
            "expires_at": expiration_time.isoformat(),
            "metadata": {
                "preparation_timestamp": datetime.utcnow().isoformat(),
                "expiration_hours": expiration_hours,
                "task_id": self.request.id
            }
        }
        
        logger.info(
            "Document delivery preparation completed",
            delivery_token=delivery_token,
            expires_at=expiration_time.isoformat()
        )
        
        return results
        
    except Exception as exc:
        logger.error("Document delivery preparation failed", error=str(exc), exc_info=True)
        
        return {
            "status": "failed",
            "error": str(exc),
            "metadata": {
                "preparation_timestamp": datetime.utcnow().isoformat(),
                "task_id": self.request.id
            }
        }


@celery_app.task(bind=True, name="app.tasks.export_tasks.send_document_notification")
def send_document_notification(
    self,
    recipient_email: str,
    delivery_url: str,
    document_name: str,
    sender_name: Optional[str] = None
) -> Dict[str, Any]:
    """
    Send document delivery notification email.
    
    Args:
        recipient_email: Email address of recipient
        delivery_url: Secure delivery URL
        document_name: Name of the document
        sender_name: Name of the sender
        
    Returns:
        Dict: Notification sending results
    """
    try:
        logger.info(
            "Sending document notification",
            recipient_email=recipient_email,
            document_name=document_name
        )
        
        # Prepare email content
        subject = f"Document Ready for Download: {document_name}"
        
        body = f"""
        Hello,
        
        A document has been prepared for you: {document_name}
        
        You can download it using the secure link below:
        {delivery_url}
        
        This link will expire in 24 hours for security purposes.
        
        {"Best regards," if not sender_name else f"Best regards,\n{sender_name}"}
        
        RealtorAgentAI Platform
        """
        
        # In production, would integrate with actual email service
        # For now, simulate email sending
        notification_result = {
            "status": "sent",
            "recipient_email": recipient_email,
            "subject": subject,
            "sent_at": datetime.utcnow().isoformat(),
            "message_id": f"msg_{self.request.id}_{datetime.utcnow().timestamp()}"
        }
        
        logger.info(
            "Document notification sent",
            recipient_email=recipient_email,
            message_id=notification_result["message_id"]
        )
        
        return notification_result
        
    except Exception as exc:
        logger.error("Document notification failed", error=str(exc), exc_info=True)
        
        return {
            "status": "failed",
            "recipient_email": recipient_email,
            "error": str(exc),
            "sent_at": datetime.utcnow().isoformat()
        }
