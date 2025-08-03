"""
Ingest tasks for file upload processing and validation.

This module contains Celery tasks for handling file uploads, document validation,
metadata extraction, and security scanning.
"""

import logging
import hashlib
import mimetypes
from typing import Dict, Any, Optional, List
from datetime import datetime
from pathlib import Path

from celery import current_task
from celery.exceptions import Retry, WorkerLostError
import structlog

from ..core.celery_app import celery_app, DatabaseTask
from ..core.storage import get_storage_client, StorageError
from ..core.document_processor import get_document_processor, DocumentProcessingError
from ..models.file import File, FileStatus, ProcessingStatus, FileProcessingJob
from ..models.audit_log import AuditLog, AuditAction
from ..services.file_service import get_file_service

logger = structlog.get_logger(__name__)


@celery_app.task(bind=True, base=DatabaseTask, name="app.tasks.ingest_tasks.process_file_upload")
def process_file_upload(
    self,
    file_id: int,
    user_id: int,
    storage_key: str,
    original_filename: str,
    file_size: int,
    mime_type: str
) -> Dict[str, Any]:
    """
    Process uploaded file with validation and metadata extraction.
    
    Args:
        file_id: Database file record ID
        user_id: User who uploaded the file
        storage_key: Storage location key
        original_filename: Original filename
        file_size: File size in bytes
        mime_type: MIME type of the file
        
    Returns:
        Dict: Processing results with status and metadata
    """
    try:
        logger.info(
            "Starting file upload processing",
            file_id=file_id,
            user_id=user_id,
            filename=original_filename,
            size=file_size,
            mime_type=mime_type
        )
        
        # Update file status to processing
        file_record = self.session.get(File, file_id)
        if not file_record:
            raise ValueError(f"File record not found: {file_id}")
        
        file_record.processing_status = ProcessingStatus.PROCESSING
        file_record.processing_started_at = datetime.utcnow()
        self.session.add(file_record)
        self.session.commit()
        
        # Get storage client and file content
        storage_client = get_storage_client()
        file_content = storage_client.download_file(storage_key)
        
        # Validate file integrity
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Virus scan (placeholder - would integrate with actual antivirus)
        virus_scan_result = virus_scan_file.delay(file_content, original_filename)
        
        # Extract basic metadata
        metadata = {
            "original_filename": original_filename,
            "file_size": file_size,
            "mime_type": mime_type,
            "file_hash": file_hash,
            "upload_timestamp": datetime.utcnow().isoformat(),
            "processing_task_id": self.request.id
        }
        
        # Validate document format and structure
        validation_result = validate_document.delay(
            file_content, original_filename, mime_type
        )
        
        # Extract detailed metadata based on file type
        metadata_result = extract_metadata.delay(
            file_content, original_filename, mime_type
        )
        
        # Update file record with initial results
        file_record.file_hash = file_hash
        file_record.metadata = metadata
        file_record.processing_status = ProcessingStatus.COMPLETED
        file_record.processing_completed_at = datetime.utcnow()
        self.session.add(file_record)
        
        # Create audit log
        audit_log = AuditLog(
            user_id=user_id,
            actor=f"system:ingest_task:{self.request.id}",
            action=AuditAction.FILE_PROCESS,
            resource_type="file",
            resource_id=str(file_id),
            success=True,
            meta={
                "file_id": file_id,
                "original_filename": original_filename,
                "file_size": file_size,
                "file_hash": file_hash,
                "processing_duration": (
                    datetime.utcnow() - file_record.processing_started_at
                ).total_seconds()
            }
        )
        self.session.add(audit_log)
        self.session.commit()
        
        logger.info(
            "File upload processing completed",
            file_id=file_id,
            file_hash=file_hash,
            processing_duration=audit_log.meta["processing_duration"]
        )
        
        return {
            "status": "completed",
            "file_id": file_id,
            "file_hash": file_hash,
            "metadata": metadata,
            "virus_scan_task_id": virus_scan_result.id,
            "validation_task_id": validation_result.id,
            "metadata_task_id": metadata_result.id
        }
        
    except Exception as exc:
        logger.error(
            "File upload processing failed",
            file_id=file_id,
            error=str(exc),
            exc_info=True
        )
        
        # Update file record with error status
        if 'file_record' in locals():
            file_record.processing_status = ProcessingStatus.FAILED
            file_record.processing_error = str(exc)
            file_record.processing_completed_at = datetime.utcnow()
            self.session.add(file_record)
            self.session.commit()
        
        # Create error audit log
        audit_log = AuditLog(
            user_id=user_id,
            actor=f"system:ingest_task:{self.request.id}",
            action=AuditAction.FILE_PROCESS,
            resource_type="file",
            resource_id=str(file_id),
            success=False,
            error_message=str(exc),
            meta={"file_id": file_id, "original_filename": original_filename}
        )
        self.session.add(audit_log)
        self.session.commit()
        
        # Retry with exponential backoff
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@celery_app.task(bind=True, name="app.tasks.ingest_tasks.validate_document")
def validate_document(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
    """
    Validate document format and structure.
    
    Args:
        file_content: File content bytes
        filename: Original filename
        mime_type: MIME type
        
    Returns:
        Dict: Validation results
    """
    try:
        logger.info("Starting document validation", filename=filename, mime_type=mime_type)
        
        validation_results = {
            "is_valid": True,
            "format_valid": True,
            "structure_valid": True,
            "errors": [],
            "warnings": []
        }
        
        # Validate MIME type matches file extension
        expected_mime = mimetypes.guess_type(filename)[0]
        if expected_mime and expected_mime != mime_type:
            validation_results["warnings"].append(
                f"MIME type mismatch: expected {expected_mime}, got {mime_type}"
            )
        
        # Validate file is not corrupted
        if len(file_content) == 0:
            validation_results["is_valid"] = False
            validation_results["errors"].append("File is empty")
        
        # Format-specific validation
        if mime_type == "application/pdf":
            validation_results.update(self._validate_pdf(file_content))
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            validation_results.update(self._validate_docx(file_content))
        elif mime_type.startswith("image/"):
            validation_results.update(self._validate_image(file_content))
        
        logger.info(
            "Document validation completed",
            filename=filename,
            is_valid=validation_results["is_valid"],
            errors=len(validation_results["errors"]),
            warnings=len(validation_results["warnings"])
        )
        
        return validation_results
        
    except Exception as exc:
        logger.error("Document validation failed", filename=filename, error=str(exc))
        return {
            "is_valid": False,
            "format_valid": False,
            "structure_valid": False,
            "errors": [f"Validation failed: {str(exc)}"],
            "warnings": []
        }

    def _validate_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Validate PDF file format."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            results = {
                "format_valid": True,
                "structure_valid": True,
                "page_count": len(doc),
                "has_text": False,
                "has_images": False
            }
            
            # Check if PDF has extractable text
            for page in doc:
                if page.get_text().strip():
                    results["has_text"] = True
                if page.get_images():
                    results["has_images"] = True
                if results["has_text"] and results["has_images"]:
                    break
            
            doc.close()
            return results
            
        except Exception as e:
            return {
                "format_valid": False,
                "structure_valid": False,
                "errors": [f"PDF validation failed: {str(e)}"]
            }

    def _validate_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Validate DOCX file format."""
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(file_content))
            
            return {
                "format_valid": True,
                "structure_valid": True,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "has_text": any(p.text.strip() for p in doc.paragraphs)
            }
            
        except Exception as e:
            return {
                "format_valid": False,
                "structure_valid": False,
                "errors": [f"DOCX validation failed: {str(e)}"]
            }

    def _validate_image(self, file_content: bytes) -> Dict[str, Any]:
        """Validate image file format."""
        try:
            from PIL import Image
            from io import BytesIO
            
            img = Image.open(BytesIO(file_content))
            
            return {
                "format_valid": True,
                "structure_valid": True,
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode
            }
            
        except Exception as e:
            return {
                "format_valid": False,
                "structure_valid": False,
                "errors": [f"Image validation failed: {str(e)}"]
            }


@celery_app.task(bind=True, name="app.tasks.ingest_tasks.extract_metadata")
def extract_metadata(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
    """
    Extract detailed metadata from uploaded file.
    
    Args:
        file_content: File content bytes
        filename: Original filename
        mime_type: MIME type
        
    Returns:
        Dict: Extracted metadata
    """
    try:
        logger.info("Starting metadata extraction", filename=filename, mime_type=mime_type)
        
        metadata = {
            "filename": filename,
            "mime_type": mime_type,
            "size": len(file_content),
            "extraction_timestamp": datetime.utcnow().isoformat()
        }
        
        # Use document processor for advanced metadata extraction
        doc_processor = get_document_processor()
        
        if mime_type in doc_processor.supported_formats:
            processing_result = await doc_processor.process_document(
                file_content, filename, mime_type.split('/')[-1]
            )
            metadata.update(processing_result.get("metadata", {}))
        
        logger.info("Metadata extraction completed", filename=filename, metadata_keys=list(metadata.keys()))
        
        return metadata
        
    except Exception as exc:
        logger.error("Metadata extraction failed", filename=filename, error=str(exc))
        return {
            "filename": filename,
            "mime_type": mime_type,
            "size": len(file_content),
            "extraction_error": str(exc),
            "extraction_timestamp": datetime.utcnow().isoformat()
        }


@celery_app.task(bind=True, name="app.tasks.ingest_tasks.virus_scan_file")
def virus_scan_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    Perform virus scan on uploaded file.
    
    Args:
        file_content: File content bytes
        filename: Original filename
        
    Returns:
        Dict: Virus scan results
    """
    try:
        logger.info("Starting virus scan", filename=filename, size=len(file_content))
        
        # Placeholder for actual virus scanning integration
        # In production, this would integrate with ClamAV, VirusTotal, or similar
        scan_results = {
            "is_clean": True,
            "threats_found": [],
            "scan_engine": "placeholder",
            "scan_timestamp": datetime.utcnow().isoformat(),
            "scan_duration": 0.1  # Placeholder duration
        }
        
        # Basic checks for suspicious content
        suspicious_patterns = [
            b"<script",
            b"javascript:",
            b"vbscript:",
            b"onload=",
            b"onerror="
        ]
        
        for pattern in suspicious_patterns:
            if pattern in file_content.lower():
                scan_results["is_clean"] = False
                scan_results["threats_found"].append(f"Suspicious pattern: {pattern.decode()}")
        
        logger.info(
            "Virus scan completed",
            filename=filename,
            is_clean=scan_results["is_clean"],
            threats=len(scan_results["threats_found"])
        )
        
        return scan_results
        
    except Exception as exc:
        logger.error("Virus scan failed", filename=filename, error=str(exc))
        return {
            "is_clean": False,
            "threats_found": [f"Scan error: {str(exc)}"],
            "scan_engine": "placeholder",
            "scan_timestamp": datetime.utcnow().isoformat(),
            "scan_error": str(exc)
        }
