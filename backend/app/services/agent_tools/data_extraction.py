"""
Data Extraction Agent Tools.

This module provides specialized tools for the Data Extraction Agent,
including document parsing, entity recognition, and confidence scoring.
"""

import re
import json
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from pathlib import Path

import structlog
from pydantic import BaseModel, Field

from .base import DataExtractionTool, ToolInput, ToolResult, ToolCategory
from ...core.config import get_settings

logger = structlog.get_logger(__name__)
settings = get_settings()


class DocumentParsingInput(ToolInput):
    """Input for document parsing tool."""
    document_id: str = Field(..., description="ID of the document to parse")
    document_type: Optional[str] = Field(None, description="Type of document (pdf, docx, image)")
    extraction_options: Dict[str, Any] = Field(default_factory=dict, description="Parsing options")


class EntityRecognitionInput(ToolInput):
    """Input for entity recognition tool."""
    text_content: str = Field(..., description="Text content to analyze")
    document_type: str = Field(default="contract", description="Type of document")
    entity_types: List[str] = Field(default_factory=list, description="Specific entity types to extract")


class ConfidenceScoringInput(ToolInput):
    """Input for confidence scoring tool."""
    extracted_entities: Dict[str, Any] = Field(..., description="Extracted entities to score")
    validation_results: Dict[str, Any] = Field(default_factory=dict, description="Validation results")
    source_metadata: Dict[str, Any] = Field(default_factory=dict, description="Source document metadata")


class DocumentParsingTool(DataExtractionTool):
    """Tool for parsing various document formats."""

    @property
    def name(self) -> str:
        return "document_parser"

    @property
    def description(self) -> str:
        return "Parse documents (PDF, DOCX, images) and extract text content with metadata"

    async def execute(self, input_data: DocumentParsingInput) -> ToolResult:
        """Parse a document and extract text content."""
        try:
            # Validate document access
            if not await self._validate_document_access(input_data.document_id, input_data.user_id):
                return ToolResult(
                    success=False,
                    errors=["Access denied to document"],
                    execution_time=0.0,
                    tool_name=self.name
                )

            # Get document metadata
            doc_metadata = await self._get_document_metadata(input_data.document_id)

            # Simulate document parsing (in real implementation, this would use
            # libraries like PyPDF2, python-docx, or OCR services)
            parsed_content = await self._parse_document_content(
                input_data.document_id,
                input_data.document_type,
                input_data.extraction_options
            )

            return ToolResult(
                success=True,
                data={
                    "document_id": input_data.document_id,
                    "text_content": parsed_content["text"],
                    "page_count": parsed_content.get("page_count", 1),
                    "word_count": len(parsed_content["text"].split()),
                    "character_count": len(parsed_content["text"]),
                    "extracted_images": parsed_content.get("images", []),
                    "tables": parsed_content.get("tables", [])
                },
                metadata={
                    "document_metadata": doc_metadata,
                    "parsing_options": input_data.extraction_options,
                    "content_type": parsed_content.get("content_type", "text")
                },
                execution_time=0.0,
                tool_name=self.name
            )

        except Exception as e:
            return ToolResult(
                success=False,
                errors=[f"Document parsing failed: {str(e)}"],
                execution_time=0.0,
                tool_name=self.name
            )

    async def _parse_document_content(self,
                                    document_id: str,
                                    document_type: Optional[str],
                                    options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse document content based on type."""
        # This is a simplified implementation
        # In production, this would integrate with actual document parsing libraries

        if document_type == "pdf":
            return await self._parse_pdf(document_id, options)
        elif document_type == "docx":
            return await self._parse_docx(document_id, options)
        elif document_type in ["jpg", "jpeg", "png", "tiff"]:
            return await self._parse_image_ocr(document_id, options)
        else:
            # Default text extraction
            return {
                "text": f"Sample extracted text from document {document_id}",
                "content_type": "text",
                "page_count": 1
            }

    async def _parse_pdf(self, document_id: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse PDF document using real PDF parsing libraries."""
        try:
            # Get document file path from database
            doc_path = await self._get_document_file_path(document_id)

            if not doc_path or not Path(doc_path).exists():
                raise FileNotFoundError(f"Document file not found: {document_id}")

            # Use PyPDF2 or similar library for actual PDF parsing
            # For now, we'll use a basic text extraction approach
            import PyPDF2

            text_content = ""
            page_count = 0

            with open(doc_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)

                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"

            return {
                "text": text_content.strip(),
                "content_type": "pdf",
                "page_count": page_count,
                "images": [],  # Would be populated by image extraction
                "tables": []   # Would be populated by table detection
            }

        except Exception as e:
            logger.error(f"PDF parsing failed for document {document_id}: {e}")
            # Fallback to basic text extraction
            return {
                "text": f"Error parsing PDF document {document_id}: {str(e)}. Please ensure the document is accessible and not corrupted.",
                "content_type": "pdf",
                "page_count": 1,
                "images": [],
                "tables": [],
                "error": str(e)
            }

    async def _parse_docx(self, document_id: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse DOCX document using real DOCX parsing libraries."""
        try:
            # Get document file path from database
            doc_path = await self._get_document_file_path(document_id)

            if not doc_path or not Path(doc_path).exists():
                raise FileNotFoundError(f"Document file not found: {document_id}")

            # Use python-docx library for actual DOCX parsing
            from docx import Document

            doc = Document(doc_path)
            text_content = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract text from tables
            tables_data = []
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text += " | ".join(row_text) + "\n"
                tables_data.append(table_text)
                text_content += table_text + "\n"

            return {
                "text": text_content.strip(),
                "content_type": "docx",
                "page_count": len(doc.paragraphs) // 20 + 1,  # Estimate pages
                "images": [],  # Would be populated by image extraction
                "tables": tables_data
            }

        except Exception as e:
            logger.error(f"DOCX parsing failed for document {document_id}: {e}")
            # Fallback to basic text extraction
            return {
                "text": f"Error parsing DOCX document {document_id}: {str(e)}. Please ensure the document is accessible and not corrupted.",
                "content_type": "docx",
                "page_count": 1,
                "images": [],
                "tables": [],
                "error": str(e)
            }

    async def _parse_image_ocr(self, document_id: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse image using OCR."""
        try:
            # Get document file path from database
            doc_path = await self._get_document_file_path(document_id)

            if not doc_path or not Path(doc_path).exists():
                raise FileNotFoundError(f"Document file not found: {document_id}")

            # Use Tesseract OCR for actual text extraction
            # Note: This requires tesseract to be installed on the system
            import pytesseract
            from PIL import Image

            # Open and process the image
            image = Image.open(doc_path)

            # Extract text using OCR
            extracted_text = pytesseract.image_to_string(image)

            # Get confidence score
            confidence_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in confidence_data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

            return {
                "text": extracted_text.strip(),
                "content_type": "image_ocr",
                "page_count": 1,
                "confidence_score": avg_confidence / 100.0  # Convert to 0-1 scale
            }

        except Exception as e:
            logger.error(f"OCR processing failed for document {document_id}: {e}")
            # Fallback response
            return {
                "text": f"Error processing image document {document_id}: {str(e)}. OCR processing failed. Please ensure Tesseract is installed and the image is readable.",
                "content_type": "image_ocr",
                "page_count": 1,
                "confidence_score": 0.0,
                "error": str(e)
            }

    async def _get_document_file_path(self, document_id: str) -> Optional[str]:
        """Get the file path for a document from the database."""
        try:
            # This would integrate with the existing file service
            # For now, we'll use a placeholder that checks common locations
            from ...services.file_service import get_file_service

            file_service = get_file_service()
            file_info = await file_service.get_file_info(document_id)

            if file_info and file_info.get("file_path"):
                return file_info["file_path"]

            # Fallback: check if it's in the uploads directory
            uploads_dir = Path("uploads")
            possible_paths = [
                uploads_dir / f"{document_id}.pdf",
                uploads_dir / f"{document_id}.docx",
                uploads_dir / f"{document_id}.jpg",
                uploads_dir / f"{document_id}.png",
            ]

            for path in possible_paths:
                if path.exists():
                    return str(path)

            return None

        except Exception as e:
            logger.error(f"Failed to get document file path for {document_id}: {e}")
            return None


class EntityRecognitionTool(DataExtractionTool):
    """Tool for recognizing real estate entities in text."""

    @property
    def name(self) -> str:
        return "entity_recognizer"

    @property
    def description(self) -> str:
        return "Extract real estate entities (properties, parties, financial terms) from text content"

    async def execute(self, input_data: EntityRecognitionInput) -> ToolResult:
        """Extract entities from text content."""
        try:
            # Extract different types of entities
            entities = await self._extract_all_entities(
                input_data.text_content,
                input_data.document_type,
                input_data.entity_types
            )

            # Validate extracted entities
            validation_results = await self._validate_entities(entities)

            return ToolResult(
                success=True,
                data={
                    "entities": entities,
                    "entity_count": sum(len(entities[category]) for category in entities),
                    "categories": list(entities.keys()),
                    "validation_results": validation_results
                },
                metadata={
                    "text_length": len(input_data.text_content),
                    "document_type": input_data.document_type,
                    "extraction_timestamp": datetime.utcnow().isoformat()
                },
                execution_time=0.0,
                tool_name=self.name
            )

        except Exception as e:
            return ToolResult(
                success=False,
                errors=[f"Entity recognition failed: {str(e)}"],
                execution_time=0.0,
                tool_name=self.name
            )

    async def _extract_all_entities(self,
                                  text: str,
                                  document_type: str,
                                  entity_types: List[str]) -> Dict[str, List[Dict[str, Any]]]:
        """Extract all types of real estate entities."""
        entities = {
            "properties": await self._extract_property_entities(text),
            "parties": await self._extract_party_entities(text),
            "financial_terms": await self._extract_financial_entities(text),
            "dates": await self._extract_date_entities(text),
            "legal_terms": await self._extract_legal_entities(text),
            "addresses": await self._extract_address_entities(text)
        }

        # Filter by requested entity types if specified
        if entity_types:
            entities = {k: v for k, v in entities.items() if k in entity_types}

        return entities

    async def _extract_property_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract property-related entities."""
        properties = []

        # Property type patterns
        property_patterns = [
            r'(?i)(single[- ]family|multi[- ]family|condominium|townhouse|apartment|commercial|residential)',
            r'(?i)(\d+[- ]bedroom|\d+[- ]bath|\d+[- ]story)',
            r'(?i)(\d+[,\d]*\s*square\s*feet|\d+[,\d]*\s*sq\.?\s*ft\.?)'
        ]

        for pattern in property_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                properties.append({
                    "type": "property_feature",
                    "value": match.group(1),
                    "start_pos": match.start(),
                    "end_pos": match.end(),
                    "confidence": 0.8
                })

        return properties

    async def _extract_party_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract party-related entities (buyers, sellers, agents)."""
        parties = []

        # Name patterns (simplified)
        name_patterns = [
            r'(?i)(buyer|seller|purchaser|vendor|agent|broker):\s*([A-Z][a-z]+\s+[A-Z][a-z]+)',
            r'(?i)([A-Z][a-z]+\s+[A-Z][a-z]+),?\s*(buyer|seller|purchaser|vendor)'
        ]

        for pattern in name_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                parties.append({
                    "type": "party",
                    "name": match.group(2) if len(match.groups()) > 1 else match.group(1),
                    "role": match.group(1) if len(match.groups()) > 1 else "unknown",
                    "start_pos": match.start(),
                    "end_pos": match.end(),
                    "confidence": 0.7
                })

        return parties

    async def _extract_financial_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract financial terms and amounts."""
        financial = []

        # Money patterns
        money_patterns = [
            r'\$[\d,]+\.?\d*',
            r'(?i)(purchase\s+price|down\s+payment|earnest\s+money|closing\s+costs):\s*\$?([\d,]+\.?\d*)',
            r'(?i)(\d+\.?\d*)\s*percent|(\d+\.?\d*)%'
        ]

        for pattern in money_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                financial.append({
                    "type": "financial_amount",
                    "value": match.group(0),
                    "start_pos": match.start(),
                    "end_pos": match.end(),
                    "confidence": 0.9
                })

        return financial

    async def _extract_date_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract date-related entities."""
        dates = []

        # Date patterns
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{1,2}-\d{1,2}-\d{4}',
            r'(?i)(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}'
        ]

        for pattern in date_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                dates.append({
                    "type": "date",
                    "value": match.group(0),
                    "start_pos": match.start(),
                    "end_pos": match.end(),
                    "confidence": 0.85
                })

        return dates

    async def _extract_legal_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract legal terms and clauses."""
        legal = []

        # Legal term patterns
        legal_patterns = [
            r'(?i)(contingency|addendum|amendment|disclosure|warranty|covenant|lien|easement)',
            r'(?i)(as[- ]is|subject\s+to|provided\s+that|notwithstanding)'
        ]

        for pattern in legal_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                legal.append({
                    "type": "legal_term",
                    "value": match.group(0),
                    "start_pos": match.start(),
                    "end_pos": match.end(),
                    "confidence": 0.75
                })

        return legal

    async def _extract_address_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract address entities."""
        addresses = []

        # Address patterns (simplified)
        address_patterns = [
            r'\d+\s+[A-Z][a-z]+\s+(Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd)',
            r'[A-Z][a-z]+,\s*[A-Z]{2}\s+\d{5}(-\d{4})?'
        ]

        for pattern in address_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                addresses.append({
                    "type": "address",
                    "value": match.group(0),
                    "start_pos": match.start(),
                    "end_pos": match.end(),
                    "confidence": 0.8
                })

        return addresses

    async def _validate_entities(self, entities: Dict[str, List[Dict[str, Any]]]) -> Dict[str, Any]:
        """Validate extracted entities for consistency and accuracy."""
        validation_results = {}

        for category, entity_list in entities.items():
            category_validation = {
                "total_entities": len(entity_list),
                "high_confidence": sum(1 for e in entity_list if e.get("confidence", 0) > 0.8),
                "medium_confidence": sum(1 for e in entity_list if 0.6 <= e.get("confidence", 0) <= 0.8),
                "low_confidence": sum(1 for e in entity_list if e.get("confidence", 0) < 0.6),
                "validation_score": sum(e.get("confidence", 0) for e in entity_list) / len(entity_list) if entity_list else 0
            }
            validation_results[category] = category_validation

        return validation_results


class ConfidenceScoringTool(DataExtractionTool):
    """Tool for calculating confidence scores for extracted data."""

    @property
    def name(self) -> str:
        return "confidence_scorer"

    @property
    def description(self) -> str:
        return "Calculate confidence scores for extracted entities and data quality assessment"

    async def execute(self, input_data: ConfidenceScoringInput) -> ToolResult:
        """Calculate confidence scores for extracted entities."""
        try:
            # Calculate overall confidence score
            overall_confidence = self._calculate_confidence_score(
                input_data.extracted_entities,
                input_data.validation_results
            )

            # Calculate category-specific scores
            category_scores = await self._calculate_category_scores(
                input_data.extracted_entities,
                input_data.validation_results
            )

            # Generate quality assessment
            quality_assessment = await self._assess_data_quality(
                input_data.extracted_entities,
                input_data.source_metadata
            )

            return ToolResult(
                success=True,
                data={
                    "overall_confidence": overall_confidence,
                    "category_scores": category_scores,
                    "quality_assessment": quality_assessment,
                    "recommendations": await self._generate_recommendations(
                        overall_confidence, category_scores, quality_assessment
                    )
                },
                metadata={
                    "scoring_timestamp": datetime.utcnow().isoformat(),
                    "entity_count": sum(len(entities) if isinstance(entities, list) else 1
                                      for entities in input_data.extracted_entities.values()),
                    "validation_available": bool(input_data.validation_results)
                },
                execution_time=0.0,
                tool_name=self.name
            )

        except Exception as e:
            return ToolResult(
                success=False,
                errors=[f"Confidence scoring failed: {str(e)}"],
                execution_time=0.0,
                tool_name=self.name
            )

    async def _calculate_category_scores(self,
                                       entities: Dict[str, Any],
                                       validation_results: Dict[str, Any]) -> Dict[str, float]:
        """Calculate confidence scores for each entity category."""
        category_scores = {}

        for category, entity_data in entities.items():
            if isinstance(entity_data, list):
                if entity_data:
                    avg_confidence = sum(e.get("confidence", 0.5) for e in entity_data) / len(entity_data)
                else:
                    avg_confidence = 0.0
            else:
                avg_confidence = 0.5  # Default for non-list data

            # Adjust based on validation results
            if category in validation_results:
                validation_score = validation_results[category].get("validation_score", 0.5)
                avg_confidence = (avg_confidence + validation_score) / 2

            category_scores[category] = round(avg_confidence, 3)

        return category_scores

    async def _assess_data_quality(self,
                                 entities: Dict[str, Any],
                                 source_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Assess overall data quality."""
        total_entities = sum(len(e) if isinstance(e, list) else 1 for e in entities.values())

        quality_metrics = {
            "completeness": min(total_entities / 10, 1.0),  # Assume 10 entities is "complete"
            "consistency": 0.8,  # Would be calculated based on cross-validation
            "accuracy": 0.75,    # Would be calculated based on validation rules
            "timeliness": 1.0    # Assume current extraction is timely
        }

        overall_quality = sum(quality_metrics.values()) / len(quality_metrics)

        return {
            "metrics": quality_metrics,
            "overall_quality": round(overall_quality, 3),
            "quality_grade": self._get_quality_grade(overall_quality)
        }

    def _get_quality_grade(self, score: float) -> str:
        """Convert quality score to letter grade."""
        if score >= 0.9:
            return "A"
        elif score >= 0.8:
            return "B"
        elif score >= 0.7:
            return "C"
        elif score >= 0.6:
            return "D"
        else:
            return "F"

    async def _generate_recommendations(self,
                                      overall_confidence: float,
                                      category_scores: Dict[str, float],
                                      quality_assessment: Dict[str, Any]) -> List[str]:
        """Generate recommendations for improving data quality."""
        recommendations = []

        if overall_confidence < 0.7:
            recommendations.append("Overall confidence is low. Consider manual review of extracted data.")

        for category, score in category_scores.items():
            if score < 0.6:
                recommendations.append(f"Low confidence in {category} extraction. Manual verification recommended.")

        if quality_assessment["overall_quality"] < 0.7:
            recommendations.append("Data quality is below acceptable threshold. Consider re-processing with different parameters.")

        if not recommendations:
            recommendations.append("Data extraction quality is acceptable. Proceed with confidence.")

        return recommendations
